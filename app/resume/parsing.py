"""Resume Parsing Engine.

Hybrid approach per spec:
  1. Raw text extraction (PDF via pypdf, DOCX via python-docx)
  2. Regex extraction for structural fields (email, phone, section headers)
  3. LLM extraction for context (skills, experience, education, projects) —
     delegated to app.resume.ai_engine, which calls this module first for
     raw_text and light regex hints, then asks the AI provider to produce
     validated structured JSON.
"""

import io
import logging
import re

from app.resume.exceptions import ResumeParsingFailedError
from app.resume.models import ResumeFileTypeEnum

logger = logging.getLogger("app.resume.parsing")

_EMAIL_RE = re.compile(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+")
_PHONE_RE = re.compile(r"(?:\+?\d{1,3}[-.\s]?)?(?:\(?\d{2,4}\)?[-.\s]?)?\d{3,4}[-.\s]?\d{3,4}")
_LINKEDIN_RE = re.compile(r"(https?://)?(www\.)?linkedin\.com/in/[A-Za-z0-9\-_/]+", re.IGNORECASE)
_GITHUB_RE = re.compile(r"(https?://)?(www\.)?github\.com/[A-Za-z0-9\-_/]+", re.IGNORECASE)

_SECTION_HEADERS = [
    "experience", "work experience", "employment", "education", "projects",
    "skills", "certifications", "achievements", "summary", "objective",
]


def extract_raw_text(file_bytes: bytes, file_type: ResumeFileTypeEnum) -> str:
    try:
        if file_type == ResumeFileTypeEnum.PDF:
            return _extract_pdf_text(file_bytes)
        elif file_type == ResumeFileTypeEnum.DOCX:
            return _extract_docx_text(file_bytes)
        raise ResumeParsingFailedError(f"Unsupported file type: {file_type}")
    except ResumeParsingFailedError:
        raise
    except Exception as exc:  # noqa: BLE001
        logger.exception("raw_text_extraction_failed")
        raise ResumeParsingFailedError(str(exc)) from exc


def _extract_pdf_text(file_bytes: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(file_bytes))
    pages = [page.extract_text() or "" for page in reader.pages]
    text = "\n".join(pages).strip()
    if not text:
        raise ResumeParsingFailedError("No extractable text found in PDF (possibly scanned/image-based).")
    return _normalize_text(text)


def _extract_docx_text(file_bytes: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    text = "\n".join(parts).strip()
    if not text:
        raise ResumeParsingFailedError("No extractable text found in DOCX.")
    return _normalize_text(text)


def _normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def regex_extract_hints(raw_text: str) -> dict:
    """Fast structural extraction used both as a fallback (if AI providers
    are unavailable) and as grounding context passed to the AI extractor to
    reduce hallucination risk.
    """
    emails = _EMAIL_RE.findall(raw_text)
    phones = _PHONE_RE.findall(raw_text)
    linkedin = _LINKEDIN_RE.search(raw_text)
    github = _GITHUB_RE.search(raw_text)

    found_sections = [h for h in _SECTION_HEADERS if h in raw_text.lower()]

    # crude name guess: first non-empty line that isn't an email/phone/url
    name_guess = None
    for line in raw_text.split("\n")[:5]:
        stripped = line.strip()
        if stripped and not _EMAIL_RE.search(stripped) and not _LINKEDIN_RE.search(stripped) and len(stripped) < 60:
            name_guess = stripped
            break

    return {
        "name_guess": name_guess,
        "email": emails[0] if emails else None,
        "phone": phones[0].strip() if phones else None,
        "linkedin_url": linkedin.group(0) if linkedin else None,
        "github_url": github.group(0) if github else None,
        "detected_sections": found_sections,
    }
